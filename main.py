# BULK LYRICS by MW DIGITAL DEVELOPMENT

# This program takes a list of songs from the user, searches for the lyrics
# to each song on Google and puts all those lyrics in a single .docx file.
# If a song's lyrics are not found, a link to the first google
# hit for that song's lyrics is saved and displayed in the document.

import re
import os
import threading
import tkinter as tk
from tkinter import messagebox, filedialog, StringVar

import requests
from bs4 import BeautifulSoup, ResultSet
from docx import Document
from docx.shared import RGBColor, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import helpers


class Application:
    def __init__(self) -> None:
        self.root = tk.Tk()
        self.root.title("Bulk Lyrics")
        self.root.geometry("600x600")

        placeholder: str = (
            "The Beatles Hey Jude\n1+1 Beyoncé"
        )

        self.textbox = tk.Text(self.root, height=20, width=50, font=("", 12), undo=True)
        self.textbox.insert(1.0, placeholder)
        self.textbox.bind("<Tab>", self.focus_next_widget)
        self.textbox.pack(pady=(40, 20))

        self.run_btn = tk.Button(self.root, text="Generate document", command=self.run)
        self.run_btn.bind("<Return>", self.run)
        self.run_btn.bind("<Tab>", self.focus_next_widget)
        self.run_btn.pack()

        self.save_btn = tk.Button(self.root, text="Save as..", command=self.save_as)
        self.save_btn.bind("<Return>", self.save_as)
        self.save_btn.bind("<Tab>", self.focus_next_widget)

        self.no_save_btn = tk.Button(
            self.root, text="Don't save", command=self.display_reset
        )
        self.no_save_btn.bind("<Return>", self.display_reset)
        self.no_save_btn.bind("<Tab>", self.focus_next_widget)

        self.status_text = StringVar()
        self.status_display = tk.Label(self.root, textvariable=self.status_text)

        self.root.mainloop()

    def run(self, *event) -> None:
        """
        Main function. Is called when the user clicks 'generate document'.
        """
        if not self.check_for_input():
            return

        self.display_running()
        self.setup_document()
        self.generate_document()
        self.display_finished()

    def check_for_input(self) -> bool:
        """
        Returns False and shows a warning when no input is found,
        Returns True if any user input is found.
        """
        if self.textbox.get("1.0", tk.END) == "\n":
            messagebox.showwarning(
                title="No Songs", message="Please enter song information."
            )
            return False
        return True

    def display_running(self) -> None:
        """
        Shows/hides relevant buttons/text when the program is running
        """
        self.run_btn.pack_forget()
        self.status_display.pack()
        self.update_status_display("Loading...\n")

    def display_reset(self) -> None:
        """
        Shows/hides relevant buttons/text when the program is reset
        """
        self.status_display.pack_forget()
        self.save_btn.pack_forget()
        self.no_save_btn.pack_forget()
        self.run_btn.pack()

    def display_finished(self) -> None:
        """
        Shows/hides relevant buttons/text when the program is finished
        """
        self.save_btn.pack(pady=10)
        self.no_save_btn.pack()
        self.update_status_display(f"100% completed")

    def update_status_display(self, text: str) -> None:
        """
        Takes a string 'text' and updates the UI's status display to it.
        """
        self.status_text.set(text)
        self.root.update()

    def focus_next_widget(self, event) -> str:
        """
        Allows the user to jump from one field/button to the next.
        """
        event.widget.tk_focusNext().focus()
        return "break"

    def setup_document(self) -> None:
        """
        Shows a loading text, initiates the docx and calls the format_document function
        """
        self.update_status_display("Preparing document...\n")
        self.document = Document()
        self.format_document(self.document)

    def format_document(self, document) -> None:
        """
        Formats the docx to make it look pretty
        """
        # Footer
        section = document.sections[0]
        footer = section.footer
        style = document.styles["Normal"]
        font = style.font

        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run("Bulk Lyrics by MW Digital Development")
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(120, 120, 120)

        # Margins
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1.27)
            section.bottom_margin = Cm(1.27)
            section.left_margin = Cm(1.27)
            section.right_margin = Cm(1.27)

        # Lyrics font
        style = document.styles["Normal"]
        font = style.font
        font.name = "Arial"
        font.size = Pt(12)

    def generate_document(self) -> None:
        """
        Takes the user's list of songs and loops over the songs.
        For each song, webscrapes Google for the lyrics.
        Then, adds the song and lyrics to the docx and adds a page break.
        """

        songlist: list = self.get_songlist()

        progess_percentage_per_song: float = 100 / len(songlist)
        total_progress_percentage: int = 0

        headers = self.start_request()

        for idx, song in enumerate(songlist):
            self.update_status_display(
                f"{round(total_progress_percentage)}% completed\n{song}"
            )

            soup: BeautifulSoup = self.fetch_song_soup(song, headers)
            song_data: dict = self.extract_song_data(song, soup)

            self.add_song_to_doc(song_data, self.document)

            # Add page break after every song, except the last one
            if idx != len(songlist) - 1:
                self.document.add_page_break()

            total_progress_percentage += progess_percentage_per_song

    def start_request(self):
        """
        Sets up a session for webscraping Google search results,
        Gets cookies,
        Returns the HTTP headers for use in fetch_song_soup function
        """
        # Set headers to mimic a web browser
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/88.0.4324.104 Safari/537.36",
        }

        # Create a session to handle cookies
        session = requests.Session()

        # Send an initial request to Google to get cookies
        session.get("https://www.google.com", headers=headers)

        return headers

    def get_songlist(self) -> list:
        """
        Gets and a list of songs from the user.
        Cleans up the list and returns it.
        """
        songlist: str = self.textbox.get("1.0", tk.END)
        songlist: list = songlist.replace("\r", "").replace('"', "").split("\n")
        # Remove redundant spaces and empty strings
        songlist = [re.sub(" +", " ", song).strip() for song in songlist if song]

        return songlist

    def fetch_song_soup(self, song: str, headers) -> BeautifulSoup:
        """
        Searches Google for a song's lyrics and returns a BeautifulSoup of
        the search results page's html.
        """
        query = self.replace_symbols_by_hex_value(song)
        page = requests.get(
            f"https://google.com/search?q={query} lyrics", headers=headers
        )
        html: str = page.text
        return BeautifulSoup(html, "lxml")


    def replace_symbols_by_hex_value(self, string: str) -> str:
        """
        Takes a string and replaces all special symbols by their
        hex value with % in front of it. e.g. '&' becomes '%26'.
        This makes sure the Google search won't lead to issues where
        a special symbol is used in the URL
        """
        for character in string:
            if not (character.isalnum() or character == ' '):
                hex_value: int = format(ord(character), "x")
                string = string.replace(character, f'%{hex_value}')

        return string


    def extract_song_data(self, song: str, soup: BeautifulSoup) -> dict:
        """
        Finds a song's title and lyrics in the song's BeautifulSoup and
        returns a dict with that info. If a song's lyrics are not found, the user's
        input is used for the song's title and google's first hit for the song's
        lyrics is stored in the dict
        """
        lyrics: ResultSet = soup.find_all("div", {"jsname": "U8S5sf"})

        if len(lyrics) == 0:
            title: str = song
            lyrics: bool = False
        else:
            title: str = soup.find("div", {"data-attrid": "title"}).text

        try:
            first_google_hit: str = soup.find("a", {"jsname": "UWckNb"})["href"]
        except:
            first_google_hit: bool = False

        song_data: dict = {
            "title": title,
            "lyrics": lyrics,
            "link": first_google_hit,
        }

        return song_data

    def add_song_to_doc(self, song_data: dict, document) -> None:
        """
        Adds a song's title and lyrics to the document
        """

        document.add_heading(song_data["title"].title())
        document.add_paragraph()

        if song_data["lyrics"]:
            for paragraph in song_data["lyrics"]:
                lines: ResultSet = paragraph.find_all("span", {"jsname": "YS01Ge"})
                p = document.add_paragraph()
                for line in lines:
                    p.add_run(line.text)
                    if line != lines[-1]:
                        p.add_run("\n")
        else:
            document.add_paragraph().add_run(
                "Lyrics Not Found"
            ).font.color.rgb = RGBColor(255, 0, 0)

            if song_data["link"]:
                p = document.add_paragraph()
                p.add_run(f"Try here: ")
                helpers.add_hyperlink(p, song_data["link"], song_data["link"])

    def save_as(self, *event) -> None:
        """
        Is called when the user clicks the 'Save as..' button.
        Prompts the user to choose a save location and then saves the docx there.
        When the file is saved, the user is asked whether they want to open the file.
        """
        filepath: str = self.choose_directory()

        if filepath:
            self.document.save(filepath)
            self.display_reset()
            self.ask_to_open_file(filepath)
        else:
            return

    def choose_directory(self) -> str:
        """
        Prompts the user to choose a save location for the docx file.
        Returns the chosen save location (path) as a string if the user has succeeded.
        """
        filetypes: list = [("Word-document", "*.docx")]
        path = None

        try:
            path = filedialog.asksaveasfile(
                filetypes=filetypes,
                defaultextension=filetypes,
                initialfile="Bulk Lyrics",
            )
        except PermissionError:
            messagebox.showwarning(
                title="Access Denied",
                message="Access denied.\nClose the document if it's open and try again.",
            )
            self.choose_directory()

        if path:
            return path.name
        else:
            return None

    def ask_to_open_file(self, path: str) -> None:
        """
        Asks the user to open a saved file or not.
        If yes, starts a thread to open the file.
        """
        open_file_response = messagebox.askyesno(
            title="Document saved",
            message=f"Document saved.\nDo you want to open it right now?",
        )

        if open_file_response:
            open_file_thread = threading.Thread(target=self.open_file, args=(path,))
            open_file_thread.start()

    def open_file(self, path: str):
        """
        Opens a file at path location
        """
        os.system('"' + path + '"')


if __name__ == "__main__":
    Application()
